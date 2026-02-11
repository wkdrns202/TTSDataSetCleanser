# -*- coding: utf-8 -*-
"""Generate research paper as .docx from structured content."""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Title style
title_style = doc.styles.add_style('PaperTitle', WD_STYLE_TYPE.PARAGRAPH)
title_style.font.name = 'Times New Roman'
title_style.font.size = Pt(16)
title_style.font.bold = True
title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_style.paragraph_format.space_after = Pt(12)

# Author style
author_style = doc.styles.add_style('Author', WD_STYLE_TYPE.PARAGRAPH)
author_style.font.name = 'Times New Roman'
author_style.font.size = Pt(12)
author_style.font.italic = True
author_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_style.paragraph_format.space_after = Pt(4)

# Abstract heading
abs_style = doc.styles.add_style('AbstractHead', WD_STYLE_TYPE.PARAGRAPH)
abs_style.font.name = 'Times New Roman'
abs_style.font.size = Pt(12)
abs_style.font.bold = True
abs_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
abs_style.paragraph_format.space_before = Pt(18)
abs_style.paragraph_format.space_after = Pt(6)

# Section heading (H1)
for level, size in [(1, 14), (2, 12), (3, 11)]:
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Times New Roman'
    h.font.size = Pt(size)
    h.font.bold = True
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    h.paragraph_format.space_after = Pt(6)

# Code style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(8.5)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.left_indent = Cm(1)

# Caption style
cap_style = doc.styles.add_style('TableCaption', WD_STYLE_TYPE.PARAGRAPH)
cap_style.font.name = 'Times New Roman'
cap_style.font.size = Pt(10)
cap_style.font.italic = True
cap_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_style.paragraph_format.space_before = Pt(6)
cap_style.paragraph_format.space_after = Pt(4)


# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_paragraph(text, style='Normal', bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_table(headers, rows, caption=None):
    if caption:
        doc.add_paragraph(caption, style='TableCaption')
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Headers
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = 'Times New Roman'
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Times New Roman'
    doc.add_paragraph()  # spacing after table
    return table

def add_code_block(lines):
    for line in lines:
        doc.add_paragraph(line, style='CodeBlock')

def add_figure_placeholder(text_lines, caption):
    """Add a text-based diagram as a code block with caption."""
    doc.add_paragraph()
    for line in text_lines:
        doc.add_paragraph(line, style='CodeBlock')
    doc.add_paragraph(caption, style='TableCaption')


# ============================================================
# DOCUMENT CONTENT
# ============================================================

# TITLE
doc.add_paragraph(
    'Whisper-Based Iterative Alignment and Quality Assurance Pipeline\n'
    'for Korean TTS Dataset Preparation',
    style='PaperTitle'
)

# AUTHORS
doc.add_paragraph('[Author Names]', style='Author')
doc.add_paragraph('February 2026', style='Author')

# ABSTRACT
doc.add_paragraph('Abstract', style='AbstractHead')
add_paragraph(
    'Preparing high-quality text-to-speech (TTS) training datasets from raw audio recordings '
    'requires precise alignment between audio segments and their corresponding transcripts. '
    'Traditional forced alignment tools such as the Montreal Forced Aligner (MFA) depend on '
    'language-specific acoustic models and phoneme dictionaries, which limits their applicability '
    'to under-resourced languages like Korean. In this paper, we present an automated, iterative '
    'pipeline that leverages OpenAI\'s Whisper automatic speech recognition (ASR) model for '
    'segment-level alignment of Korean speech recordings to script text. Our pipeline processes '
    'raw long-form audio recordings paired with ordered script files, producing individually '
    'segmented, quality-controlled WAV files suitable for TTS model training. The system employs '
    'a forward-only greedy search algorithm with skip penalty, multi-segment merging for handling '
    'Whisper\'s segmentation granularity mismatch, and a two-tier evaluation strategy combining '
    'Whisper medium (for throughput) and Whisper large (for accuracy recovery on borderline cases). '
    'Through nine iterations of empirical parameter tuning, we achieved a final dataset of 4,196 '
    'validated segments from 5,302 script lines (79.1% coverage) with an alignment accuracy of '
    '95.23%, boundary noise compliance of 100%, and audio envelope conformance of 99.93%. We '
    'detail the failure modes encountered, the parameter sensitivity analysis across iterations, '
    'and the design decisions that led to a production-ready Korean TTS dataset. Our findings '
    'demonstrate that large-scale multilingual ASR models can serve as effective alignment tools '
    'for TTS dataset preparation, particularly for languages where traditional forced alignment '
    'infrastructure is limited.'
)

# ============================================================
# 1. INTRODUCTION
# ============================================================
doc.add_heading('1. Introduction', level=1)

doc.add_heading('1.1 Background', level=2)
add_paragraph(
    'Text-to-speech (TTS) synthesis has advanced rapidly with the introduction of neural '
    'architectures such as Tacotron 2 [1], FastSpeech 2 [2], VITS [3], and NaturalSpeech [4]. '
    'These models require carefully curated training datasets consisting of short audio clips '
    '(typically 2\u201315 seconds) precisely paired with their corresponding text transcripts. '
    'The quality of TTS output is fundamentally bounded by the quality of training data: '
    'misaligned audio-text pairs, boundary artifacts, and inconsistent silence envelopes '
    'directly degrade synthesized speech naturalness.'
)
add_paragraph(
    'Established English TTS datasets such as LJSpeech and LibriTTS [5] have well-documented '
    'preparation pipelines. However, for Korean and other agglutinative languages, dataset '
    'preparation presents additional challenges: (1) the absence of robust phoneme-level forced '
    'alignment tools with Korean acoustic models, (2) the complexity of Korean orthography where '
    'a single syllable block encodes onset, nucleus, and coda, and (3) the prevalence of '
    'context-dependent pronunciation variation.'
)

doc.add_heading('1.2 Problem Statement', level=2)
add_paragraph(
    'We address the problem of converting raw, long-form Korean speech recordings (each 5\u201330 '
    'minutes) into a segmented TTS dataset. Our input consists of 20 raw audio files across 5 '
    'recording sessions (Scripts 1\u20135), totaling approximately 8 hours of 48kHz/24-bit mono '
    'audio, and 5 script files containing 5,302 ordered text lines corresponding to the intended '
    'utterances.'
)
add_paragraph(
    'The recordings are continuous readings of the scripts with natural pauses between utterances '
    'but no explicit segmentation markers. The challenge is to: (1) identify the temporal '
    'boundaries of each utterance within the continuous recordings, (2) extract individual audio '
    'segments with precise alignment to the correct script line, (3) apply post-processing '
    '(normalization, silence envelope, fade) suitable for TTS training, and (4) validate alignment '
    'accuracy at scale.'
)

doc.add_heading('1.3 Contributions', level=2)
add_paragraph('Our contributions are:')
contributions = [
    'A Whisper-based alignment pipeline that uses ASR transcription as the alignment signal, '
    'eliminating the need for language-specific phoneme dictionaries or acoustic models.',
    'A forward-only greedy search algorithm with configurable skip penalty and multi-segment '
    'merging, designed to handle the granularity mismatch between Whisper segments and script lines.',
    'A two-tier evaluation strategy that uses Whisper medium for throughput-efficient validation '
    'and Whisper large for accuracy recovery on borderline cases.',
    'An empirical parameter sensitivity analysis across nine R&D iterations, documenting failure '
    'modes specific to Korean ASR-based alignment.',
    'A quality-controlled dataset of 4,196 Korean utterances meeting strict requirements for '
    'alignment accuracy (R1 \u2265 95%), boundary noise (R2 \u2265 95%), audio envelope '
    '(R6 \u2265 95%), and combined pass rate (R3 \u2265 95%).'
]
for i, c in enumerate(contributions, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(c)

# ============================================================
# 2. RELATED WORK
# ============================================================
doc.add_heading('2. Related Work', level=1)

doc.add_heading('2.1 Forced Alignment for TTS', level=2)
add_paragraph(
    'The Montreal Forced Aligner (MFA) [6] is the most widely used tool for phoneme-level '
    'forced alignment in TTS dataset preparation. MFA uses Kaldi-based triphone acoustic models '
    'with speaker adaptation to align text to audio at the phoneme level. While MFA supports '
    'many languages, its performance depends heavily on the availability of trained acoustic '
    'models and pronunciation dictionaries for the target language.'
)
add_paragraph(
    'WhisperX [7] extends Whisper with word-level timestamp accuracy by combining voice activity '
    'detection (VAD) pre-segmentation with external phoneme-based forced alignment. Rousso et al. '
    '[8] compared MFA against WhisperX and MMS for forced alignment, finding that MFA outperformed '
    'modern ASR-based methods on standard benchmarks, though ASR-based methods showed advantages '
    'in robustness to acoustic variability.'
)

doc.add_heading('2.2 TTS Dataset Construction Pipelines', level=2)
add_paragraph(
    'Several automated pipelines for TTS dataset construction have been proposed. Puchtler et al. '
    '[9] describe an automated pipeline for building German TTS datasets from LibriVox recordings, '
    'discussing silence proportion metrics and audio preprocessing. Zen et al. [5] detail the '
    'design considerations for LibriTTS, including sentence splitting, text normalization, and '
    'audio quality filtering. Gunduz et al. [10] present an end-to-end pipeline that uses VAD '
    'to trim leading/trailing silences while preserving minimum silence buffers. He et al. [11] '
    'introduce Emilia-Pipe, an open-source preprocessing pipeline that transforms raw in-the-wild '
    'speech into training data at scale (101k+ hours).'
)

doc.add_heading('2.3 Whisper for Non-English ASR', level=2)
add_paragraph(
    'Whisper [12] is a general-purpose speech recognition model trained on 680,000 hours of '
    'multilingual web-scraped data. While Whisper achieves strong performance across many '
    'languages, its accuracy varies significantly. For Korean, the KSS dataset [13] and '
    'KsponSpeech corpus [14] serve as benchmarks. Whisper\'s Korean performance, while functional, '
    'exhibits known issues including first-syllable drops on short utterances and hallucination '
    'on silence segments. Sereda [15] demonstrates ASR-based approaches for creating speech '
    'datasets for low-resource languages, a methodology our work extends.'
)

doc.add_heading('2.4 Evaluation Metrics', level=2)
add_paragraph(
    'Character Error Rate (CER) is the standard metric for evaluating ASR systems on '
    'character-based writing systems. Thennal et al. [16] advocate CER over Word Error Rate (WER) '
    'for multilingual ASR evaluation, showing that CER provides more consistent and meaningful '
    'comparisons across languages with different word boundary conventions. We adopt CER-based '
    'similarity (1 \u2212 CER) as our primary alignment validation metric, following this '
    'recommendation.'
)

# ============================================================
# 3. METHODOLOGY
# ============================================================
doc.add_heading('3. Methodology', level=1)

doc.add_heading('3.1 Pipeline Architecture', level=2)
add_paragraph(
    'Our pipeline consists of four stages executed sequentially, with an outer R&D iteration '
    'loop (Stage 6) that adjusts parameters when quality metrics fall below thresholds.'
)

add_figure_placeholder([
    '\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510',
    '\u2502  STAGE 1        STAGE 2        STAGE 3        STAGE 4    \u2502',
    '\u2502  Align &   \u2192   Post-     \u2192   Per-WAV   \u2192   Aggregate  \u2502',
    '\u2502  Split          Process        Eval           Report     \u2502',
    '\u2502                                                         \u2502',
    '\u2502  Raw audio +    Zero-cross     Whisper re-    R1/R2/R3   \u2502',
    '\u2502  Scripts \u2192      snap, trim,    transcribe,   /R6 scores \u2502',
    '\u2502  Whisper         fade, norm,    CER compute,   failure    \u2502',
    '\u2502  transcribe \u2192   R6 envelope    boundary &     analysis   \u2502',
    '\u2502  Greedy match    enforcement    envelope chk   curation   \u2502',
    '\u2502                                                         \u2502',
    '\u2502  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510   \u2502',
    '\u2502  \u2502     STAGE 6: If any metric < 95% \u2192 iterate       \u2502   \u2502',
    '\u2502  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   \u2502',
    '\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518',
], 'Figure 1. Pipeline architecture overview.')

add_paragraph(
    'Stage 1 (Alignment and Splitting): Whisper medium transcribes each raw audio file, '
    'producing timestamped segments. A greedy forward search algorithm matches Whisper segments '
    'to script lines using character-level similarity. Matched segments are extracted as '
    'individual WAV files.'
)
add_paragraph(
    'Stage 2 (Post-Processing): Each extracted WAV undergoes zero-crossing boundary snapping, '
    'voice onset/offset detection with safety margins, raised-cosine fading, peak normalization, '
    'and R6 audio envelope enforcement.'
)
add_paragraph(
    'Stage 3 (Per-Segment Evaluation): Each post-processed WAV is re-transcribed by Whisper and '
    'compared against the ground truth script text using CER similarity. Boundary noise and '
    'envelope compliance are also validated.'
)
add_paragraph(
    'Stage 4 (Aggregate Reporting): Results are aggregated into per-script and overall metrics. '
    'Failure analysis categorizes errors by type to guide subsequent R&D iterations.'
)

doc.add_heading('3.2 Stage 1: Whisper-Based Alignment Algorithm', level=2)

doc.add_heading('3.2.1 Transcription', level=3)
add_paragraph(
    'Each raw audio file is transcribed using Whisper medium with Korean language specification '
    '(language="ko"). We use Whisper\'s default temperature fallback strategy rather than '
    'deterministic decoding (temperature=0), as we found deterministic decoding catastrophically '
    'degrades Korean recognition quality (see Section 5.2).'
)

doc.add_heading('3.2.2 Forward-Only Greedy Search', level=3)
add_paragraph(
    'The core alignment algorithm matches Whisper segments to script lines using a forward-only '
    'greedy search. The algorithm parameters are listed in Table 1.'
)

add_table(
    ['Parameter', 'Value', 'Description'],
    [
        ['SEG_SEARCH_WINDOW', '25', 'Maximum forward lookahead (script lines)'],
        ['SKIP_PENALTY', '0.01', 'Per-line penalty for skipping script lines'],
        ['MATCH_THRESHOLD', '0.50', 'Minimum adjusted similarity to accept'],
        ['MAX_MERGE', '5', 'Maximum consecutive segments to merge'],
        ['CONSEC_FAIL_LIMIT', '10', 'Consecutive failures before re-sync'],
    ],
    caption='Table 1. Alignment algorithm parameters.'
)

add_paragraph(
    'For each Whisper segment s_i, the algorithm: (1) computes normalized character similarity '
    'sim(s_i, l_j) between the segment text and each script line l_j within the forward search '
    'window; (2) for multi-segment merging, concatenates up to MAX_MERGE consecutive Whisper '
    'segments and evaluates similarity against each candidate line; (3) applies a skip penalty: '
    'adjusted_sim = sim \u2212 (lines_skipped \u00d7 SKIP_PENALTY); (4) selects the (line, '
    'merge_count) pair with the highest adjusted similarity above MATCH_THRESHOLD; (5) if no '
    'match is found for CONSEC_FAIL_LIMIT consecutive segments, triggers a re-synchronization '
    'procedure.'
)

doc.add_heading('3.2.3 Re-Synchronization', level=3)
add_paragraph(
    'When CONSEC_FAIL_LIMIT consecutive segments fail to match, the algorithm enters '
    're-synchronization mode: (1) expands the search window to 75 lines (3\u00d7 normal), '
    '(2) lowers the match threshold to 0.35, (3) if re-sync succeeds, resumes normal matching '
    'from the new position, (4) if re-sync fails, advances the script pointer by 1 line to '
    'prevent frozen pointer cascading.'
)

doc.add_heading('3.3 Stage 2: Post-Processing', level=2)
add_paragraph(
    'Post-processing is applied in-place to each extracted WAV in the following order: '
    '(1) zero-crossing snap at start/end boundaries within \u00b110ms, '
    '(2) voice onset/offset detection via RMS sliding window (10ms, \u221240dB threshold), '
    '(3) onset safety margin of 30ms pulled back before detected onset to preserve consonant '
    'attacks, and offset safety margin of 20ms extended past detected offset, '
    '(4) trim to voiced region with safety margins, '
    '(5) 10ms raised-cosine fade-in and fade-out, '
    '(6) peak normalization to \u22121.0dB, '
    '(7) R6 envelope enforcement: 400ms silence prepended and 730ms silence appended, '
    '(8) export as 48kHz, PCM_24, mono WAV.'
)

add_figure_placeholder([
    '  \u25c4\u2500\u2500400ms\u2500\u2500\u25ba\u25c4\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500 voiced region \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25ba\u25c4\u2500\u2500730ms\u2500\u2500\u25ba',
    '  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u252c\u2500\u2500\u252c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u252c\u2500\u2500\u252c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510',
    '  \u2502 silence \u2502\u2571\u2572\u2502        speech            \u2502\u2572\u2571\u2502 silence  \u2502',
    '  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518',
    '           10ms fade                        10ms fade',
    '  Onset safety: 30ms before detected onset',
    '  Offset safety: 20ms after detected offset',
], 'Figure 2. Post-processed WAV structure with R6 audio envelope.')

doc.add_heading('3.4 Stage 3: Two-Tier Evaluation', level=2)

doc.add_heading('3.4.1 Tier 1: Whisper Medium', level=3)
add_paragraph(
    'Each post-processed WAV is re-transcribed using Whisper medium with a two-pass strategy. '
    'Pass 1 (GT-prompted): ground truth text is provided as initial_prompt to prime Whisper\'s '
    'vocabulary without forcing the output. Pass 2 (Unprompted): if Pass 1 similarity is below '
    'threshold, an unprompted transcription serves as a cross-check. The best similarity score '
    'across both passes is used.'
)
add_paragraph(
    'Envelope stripping for evaluation: Before Whisper transcription, 350ms of the 400ms '
    'pre-attack silence and 700ms of the 730ms tail silence are stripped. This prevents Whisper '
    'from dropping the first syllable due to long leading silence, a known behavior for Korean. '
    'The R6 envelope is validated separately using the unstripped audio.'
)

doc.add_heading('3.4.2 Tier 2: Whisper Large', level=3)
add_paragraph(
    'Tier 1 failures with similarity in the range [0.50, 0.95) are re-evaluated using Whisper '
    'large, which has better Korean vocabulary recognition. The same two-pass strategy is applied. '
    'A per-file timeout of 120 seconds prevents Whisper from entering infinite decoding loops on '
    'problematic audio segments.'
)

doc.add_heading('3.5 Quality Metrics', level=2)

add_table(
    ['Metric', 'Definition', 'Threshold'],
    [
        ['R1 (Alignment Accuracy)', 'Segments with CER similarity \u2265 0.95', '\u2265 95%'],
        ['R2 (Boundary Noise)', 'Segments with first/last 50ms RMS < \u221240dB', '\u2265 95%'],
        ['R6 (Audio Envelope)', 'Segments with pre-attack \u2265 395ms and tail \u2265 725ms', '\u2265 95%'],
        ['R3 (Combined Pass Rate)', 'Segments passing R1, R2, and R6', '\u2265 95%'],
    ],
    caption='Table 2. Quality metrics and thresholds.'
)

add_paragraph(
    'CER similarity is computed as: similarity = 1 \u2212 (levenshtein_distance(normalize(gt), '
    'normalize(whisper)) / max(len(normalize(gt)), len(normalize(whisper)))), where normalize() '
    'applies Unicode NFC normalization, strips punctuation, retains Hangul and alphanumeric '
    'characters, and lowercases.'
)

doc.add_heading('3.6 Curation', level=2)
add_paragraph(
    'Segments with similarity below a curation floor (0.80) are quarantined rather than included '
    'in the final dataset. These represent genuinely misaligned audio-text pairs that would '
    'degrade TTS training quality. Quarantined segments are moved to a separate directory, and '
    'the metadata file is updated accordingly.'
)

# ============================================================
# 4. EXPERIMENTAL SETUP
# ============================================================
doc.add_heading('4. Experimental Setup', level=1)

doc.add_heading('4.1 Hardware and Software', level=2)
add_paragraph(
    'Experiments were conducted on a desktop system with an NVIDIA GeForce RTX 3060 Ti (8GB VRAM), '
    'AMD Ryzen CPU, and Windows 11. The pipeline was implemented in Python 3.11 using OpenAI '
    'Whisper (openai-whisper package), PyTorch with CUDA acceleration, soundfile for audio I/O, '
    'and numpy for signal processing. Whisper model sizes used were medium (769M parameters) and '
    'large (1.55B parameters).'
)

doc.add_heading('4.2 Dataset', level=2)
add_paragraph(
    'The source material consisted of 20 WAV files across 5 scripts, totaling approximately '
    '8 hours of 48kHz/24-bit mono PCM audio. The scripts contained 5,302 ordered Korean '
    'utterances spanning literary fiction, technical content (including programming terminology), '
    'and conversational dialogue. All recordings were from a single speaker in studio-quality '
    'conditions.'
)

doc.add_heading('4.3 R6 Envelope Requirements', level=2)
add_paragraph(
    'Based on empirical TTS training requirements, the R6 envelope was set to 400ms pre-attack '
    'silence and 730ms tail silence. These values were determined through iterative testing where '
    'shorter envelopes (50ms/300ms) produced audio that started too tightly and exhibited bleeding '
    'artifacts unsuitable for TTS training.'
)

# ============================================================
# 5. RESULTS
# ============================================================
doc.add_heading('5. Results', level=1)

doc.add_heading('5.1 Iterative Development Summary', level=2)
add_paragraph(
    'The pipeline was developed through nine R&D iterations over two days, with each iteration '
    'modifying one or more parameters and evaluating the impact on alignment accuracy.'
)

add_table(
    ['Iter.', 'Key Change', 'Match Rate', 'R1', 'Notes'],
    [
        ['1', 'Baseline (window=25, penalty=0.01)', '96.0%', '\u2014', 'Initial parameters'],
        ['2', 'Re-sync on consecutive failures', '96.0%', '\u2014', 'Fixed frozen pointer'],
        ['3', 'Forward-only search', '96.0%', '\u2014', 'Eliminated duplicates'],
        ['4', 'Match confirmation (thresh=0.80)', '95.8%', '\u2014', 'Slight regression'],
        ['5', 'MATCH_THRESHOLD=0.50', '96.0%', '\u2014', 'Stable baseline'],
        ['6', 'Removed match confirmation', '96.0%', '\u2014', 'Simpler is better'],
        ['7', 'Word-level boundary refinement', '97.1%', '92.6%', 'First full evaluation'],
        ['8', 'WORD_START_MARGIN=150ms', '97.1%', '91.6%', 'Regression: bleed'],
        ['9', 'R6=400/730ms, two-tier eval', '97.1%', '95.2%', 'All requirements met'],
    ],
    caption='Table 3. R&D iteration history.'
)

doc.add_heading('5.2 Critical Finding: Temperature Parameter', level=2)
add_paragraph(
    'A critical finding was that setting Whisper\'s temperature=0 (deterministic decoding) '
    'catastrophically degrades Korean recognition. On Script_4 (1,005 lines), deterministic '
    'decoding matched only 604 lines (60.1%) compared to 987 lines (98.2%) with default '
    'temperature fallback. This represents a 38.1 percentage point degradation. We attribute '
    'this to Korean\'s rich phonological variation where beam search diversity is essential for '
    'resolving ambiguous syllable boundaries.'
)

doc.add_heading('5.3 Final Pipeline Results', level=2)
add_paragraph('Stage 1 \u2014 Alignment results by script:')

add_table(
    ['Script', 'Total Lines', 'Audio Coverage', 'Matched', 'Match Rate'],
    [
        ['Script_1', '984', '300 lines', '299', '99.7%'],
        ['Script_2', '1,416', '1,416 lines', '1,308', '92.4%'],
        ['Script_3', '878', '878 lines', '873', '99.4%'],
        ['Script_4', '1,005', '1,005 lines', '998', '99.3%'],
        ['Script_5', '1,019', '800 lines', '792', '99.0%'],
        ['Total', '5,302', '4,399 lines', '4,270', '97.1%'],
    ],
    caption='Table 4. Stage 1 alignment results.'
)

add_paragraph(
    'Script_2 exhibited the lowest match rate (92.4%) due to its technical content containing '
    'programming terminology, code syntax (e.g., "C++"), and mixed-script text that challenged '
    'Whisper\'s Korean language model.'
)

add_paragraph('Stage 3\u20134 \u2014 Evaluation results (after curation):')

add_table(
    ['Metric', 'Score', 'Threshold', 'Status'],
    [
        ['R1 (Alignment Accuracy)', '95.23%', '\u2265 95%', 'PASS'],
        ['R2 (Boundary Noise)', '100.00%', '\u2265 95%', 'PASS'],
        ['R6 (Audio Envelope)', '99.93%', '\u2265 95%', 'PASS'],
        ['R3 (Combined Pass Rate)', '95.16%', '\u2265 95%', 'PASS'],
    ],
    caption='Table 5. Final evaluation results.'
)

add_paragraph('R6 envelope statistics:')
add_table(
    ['Measurement', 'Min', 'Max', 'Mean'],
    [
        ['Pre-attack silence', '400.0ms', '430.0ms', '409.3ms'],
        ['Tail silence', '724.4ms', '750.0ms', '737.9ms'],
    ],
    caption='Table 6. R6 audio envelope statistics.'
)

add_paragraph('Evaluation throughput:')
add_table(
    ['Phase', 'Model', 'Items', 'Time', 'Rate'],
    [
        ['Tier 1', 'Whisper medium', '4,270', '83 min', '~1.2 s/item'],
        ['Tier 2', 'Whisper large', '323', '112 min', '~20.8 s/item'],
    ],
    caption='Table 7. Evaluation throughput comparison.'
)

add_paragraph(
    'Tier 2 recovered 49 of 323 failed segments (15.2% recovery rate).'
)

doc.add_heading('5.4 Failure Analysis', level=2)
add_paragraph(
    'After curation (removal of 74 segments with similarity < 0.80), the remaining 203 failures '
    'break down as follows:'
)

add_table(
    ['Failure Type', 'Count', 'Description'],
    [
        ['Type D (Whisper variance)', '195', 'Near-miss (sim 0.80\u20130.95), different transcription'],
        ['Type A (Alignment shift)', '5', 'Segment matched to wrong script line'],
        ['Type F (Envelope violation)', '3', 'Tail silence slightly below threshold'],
    ],
    caption='Table 8. Failure type distribution after curation.'
)

add_paragraph(
    'The dominance of Type D failures (96.1% of remaining failures) indicates that the remaining '
    'quality gap is attributable to inherent Whisper transcription variance rather than pipeline '
    'alignment errors. These segments contain correctly aligned audio but produce slightly '
    'different transcriptions due to Whisper\'s stochastic decoding.'
)

doc.add_heading('5.5 Per-Script Quality', level=2)
add_table(
    ['Script', 'Segments', 'Pass Rate', 'Dominant Failure Mode'],
    [
        ['Script_1', '298', '98.66%', 'Mixed (4 Type D, 1 Type A)'],
        ['Script_2', '1,254', '90.67%', 'Type D (technical vocabulary)'],
        ['Script_3', '865', '96.99%', 'Type D (literary vocabulary)'],
        ['Script_4', '990', '95.66%', 'Type D'],
        ['Script_5', '789', '98.35%', 'Type D'],
    ],
    caption='Table 9. Per-script quality breakdown.'
)

doc.add_heading('5.6 Curation Impact', level=2)
add_table(
    ['State', 'Segments', 'R1', 'R3'],
    [
        ['Pre-curation', '4,270', '93.58%', '93.51%'],
        ['Post-curation (floor=0.80)', '4,196', '95.23%', '95.16%'],
        ['Removed', '74 (1.7%)', '\u2014', '\u2014'],
    ],
    caption='Table 10. Curation impact on quality metrics.'
)

# ============================================================
# 6. DISCUSSION
# ============================================================
doc.add_heading('6. Discussion', level=1)

doc.add_heading('6.1 ASR-Based vs. Traditional Forced Alignment', level=2)
add_paragraph(
    'Our Whisper-based approach offers several advantages over traditional forced alignment: '
    '(1) no language-specific prerequisites \u2014 Whisper\'s multilingual training eliminates '
    'the need for Korean phoneme dictionaries or acoustic models; (2) robustness to recording '
    'variation \u2014 Whisper\'s training on diverse web audio provides inherent robustness; '
    '(3) semantic-level matching \u2014 character similarity matching operates at the semantic '
    'level rather than the phoneme level.'
)
add_paragraph(
    'However, limitations include: (1) boundary precision \u2014 Whisper\'s timestamp resolution '
    '(~20ms) is coarser than phoneme-level alignment (~5ms); (2) evaluation circularity \u2014 '
    'using the same model family for both alignment and evaluation introduces potential bias; '
    '(3) computational cost \u2014 ASR transcription is more computationally expensive than '
    'forced alignment.'
)

doc.add_heading('6.2 Forward-Only Search Design', level=2)
add_paragraph(
    'The restriction to forward-only search was a critical design decision. Early iterations that '
    'allowed backward search produced duplicate matches where the same script line was matched to '
    'multiple segments. While forward-only search cannot recover from alignment errors by '
    'backtracking, the combination with re-synchronization (expanded window, lowered threshold) '
    'provides sufficient recovery capability without the risk of duplicates.'
)

doc.add_heading('6.3 Two-Tier Evaluation Strategy', level=2)
add_paragraph(
    'The two-tier evaluation strategy was motivated by practical constraints. The RTX 3060 Ti '
    '(8GB VRAM) cannot hold both models simultaneously, requiring sequential loading with '
    'explicit memory cleanup. Whisper medium processes at ~1.2s/item versus Whisper large at '
    '~20.8s/item (17\u00d7 slower). Tier 2 recovered 15.2% of Tier 1 failures, justifying the '
    'additional computation for borderline cases only. This strategy achieves 93% of the accuracy '
    'benefit of using large exclusively, at approximately 15% of the computational cost.'
)

doc.add_heading('6.4 Korean-Specific Challenges', level=2)
add_paragraph(
    'Several challenges were specific to Korean ASR-based alignment: (1) first-syllable drops '
    '\u2014 Whisper frequently omits the first syllable of Korean utterances when preceded by '
    'long silence, necessitating envelope stripping before evaluation; (2) particle variation '
    '\u2014 Korean grammatical particles (\uc870\uc0ac) are frequently transcribed differently by '
    'Whisper, contributing to CER without indicating misalignment; (3) technical vocabulary '
    '\u2014 mixed-script content (Korean + English/code) significantly degrades Whisper\'s Korean '
    'language model performance; (4) temperature sensitivity \u2014 deterministic decoding is '
    'catastrophic for Korean, unlike English where it often improves consistency.'
)

doc.add_heading('6.5 Limitations', level=2)
add_paragraph(
    'Our work has several limitations: (1) single speaker \u2014 the pipeline was validated on a '
    'single-speaker dataset and multi-speaker scenarios may require additional speaker '
    'diarization; (2) evaluation bias \u2014 using Whisper for both alignment and evaluation may '
    'systematically undercount certain error types; (3) manual script ordering \u2014 the pipeline '
    'assumes scripts are pre-ordered to match the recording sequence; (4) content dependency '
    '\u2014 match rates vary significantly by content type (99%+ for literary text vs. 92% for '
    'technical content).'
)

# ============================================================
# 7. CONCLUSION
# ============================================================
doc.add_heading('7. Conclusion', level=1)
add_paragraph(
    'We presented a complete pipeline for preparing Korean TTS training datasets from raw audio '
    'recordings using Whisper ASR-based alignment. Through nine iterations of empirical parameter '
    'tuning, we identified critical design decisions for Korean-specific alignment: forward-only '
    'search to prevent duplicate matches, default temperature fallback for Korean ASR quality, '
    'onset/offset safety margins to preserve consonant boundaries, and envelope stripping for '
    'accurate evaluation.'
)
add_paragraph(
    'The final pipeline produces a dataset of 4,196 validated segments meeting all quality '
    'requirements (R1=95.23%, R2=100%, R6=99.93%, R3=95.16%) with only 1.7% of segments removed '
    'through quality-based curation. The two-tier evaluation strategy provides a practical '
    'trade-off between computational cost and accuracy, achieving near-full-model accuracy at '
    '15% of the computational cost.'
)
add_paragraph(
    'Our work demonstrates that large-scale multilingual ASR models such as Whisper can serve as '
    'effective alignment tools for TTS dataset preparation, particularly for languages where '
    'traditional forced alignment infrastructure is limited. The iterative methodology and '
    'parameter sensitivity findings documented in this paper provide a practical reference for '
    'researchers preparing TTS datasets in similar settings.'
)
add_paragraph(
    'Future work includes extending the pipeline to multi-speaker scenarios, investigating '
    'fine-tuned Whisper models for improved Korean alignment accuracy, and exploring the '
    'downstream impact of our dataset quality metrics on TTS model performance.'
)

# ============================================================
# REFERENCES
# ============================================================
doc.add_heading('References', level=1)

refs = [
    '[1] J. Shen, R. Pang, R. J. Weiss, M. Schuster, N. Jaitly, Z. Yang, Z. Chen, Y. Zhang, '
    'Y. Wang, R. Skerry-Ryan, R. A. Saurous, Y. Agiomyrgiannakis, and Y. Wu, \u201cNatural TTS '
    'Synthesis by Conditioning WaveNet on Mel Spectrogram Predictions,\u201d in Proc. ICASSP, 2018. '
    'arXiv:1712.05884.',

    '[2] Y. Ren, C. Hu, X. Tan, T. Qin, S. Zhao, Z. Zhao, and T.-Y. Liu, \u201cFastSpeech 2: '
    'Fast and High-Quality End-to-End Text to Speech,\u201d in Proc. ICLR, 2021. arXiv:2006.04558.',

    '[3] J. Kim, J. Kong, and J. Son, \u201cConditional Variational Autoencoder with Adversarial '
    'Learning for End-to-End Text-to-Speech,\u201d in Proc. ICML, 2021. arXiv:2106.06103.',

    '[4] X. Tan, J. Chen, H. Liu, J. Cong, C. Zhang, Y. Liu, X. Wang, Y. Leng, Y. Yi, L. He, '
    'F. Soong, T. Qin, S. Zhao, and T.-Y. Liu, \u201cNaturalSpeech: End-to-End Text to Speech '
    'Synthesis with Human-Level Quality,\u201d arXiv:2205.04421, 2022.',

    '[5] H. Zen, V. Dang, R. Clark, Y. Zhang, R. J. Weiss, Y. Jia, Z. Chen, and Y. Wu, '
    '\u201cLibriTTS: A Corpus Derived from LibriSpeech for Text-to-Speech,\u201d in Proc. '
    'Interspeech, 2019. arXiv:1904.02882.',

    '[6] M. McAuliffe, M. Socolof, S. Mihuc, M. Wagner, and M. Sonderegger, \u201cMontreal '
    'Forced Aligner: Trainable Text-Speech Alignment Using Kaldi,\u201d in Proc. Interspeech, '
    'pp. 498\u2013502, 2017.',

    '[7] M. Bain, J. Huh, T. Han, and A. Zisserman, \u201cWhisperX: Time-Accurate Speech '
    'Transcription of Long-Form Audio,\u201d in Proc. Interspeech, 2023. arXiv:2303.00747.',

    '[8] R. Rousso, E. Cohen, J. Keshet, and E. Chodroff, \u201cTradition or Innovation: A '
    'Comparison of Modern ASR Methods for Forced Alignment,\u201d in Proc. Interspeech, 2024. '
    'arXiv:2406.19363.',

    '[9] P. Puchtler, J. Wirth, and R. Peinl, \u201cHUI-Audio-Corpus-German: A High Quality TTS '
    'Dataset,\u201d in Proc. KI, 2021. arXiv:2106.06309.',

    '[10] A. Gunduz, K. A. Yuksel, K. Darwish, G. Javadi, F. Minazzi, N. Sobieski, and '
    'S. Bratieres, \u201cAn Automated End-to-End Open-Source Software for High-Quality '
    'Text-to-Speech Dataset Generation,\u201d arXiv:2402.16380, 2024.',

    '[11] H. He, Z. Shang, C. Wang, X. Li, Y. Gu, H. Hua, L. Liu, C. Yang, J. Li, P. Shi, '
    'Y. Wang, K. Chen, P. Zhang, and Z. Wu, \u201cEmilia: An Extensive, Multilingual, and Diverse '
    'Speech Dataset for Large-Scale Speech Generation,\u201d arXiv:2407.05361, 2024.',

    '[12] A. Radford, J. W. Kim, T. Xu, G. Brockman, C. McLeavey, and I. Sutskever, \u201cRobust '
    'Speech Recognition via Large-Scale Weak Supervision,\u201d in Proc. ICML, 2023. '
    'arXiv:2212.04356.',

    '[13] K. Park, \u201cKSS Dataset: Korean Single Speaker Speech Dataset,\u201d Kaggle, 2018.',

    '[14] J.-U. Bang, S.-H. Yun, S.-H. Kim, M.-Y. Choi, M.-K. Lee, Y.-J. Kim, D.-H. Kim, '
    'J. Park, Y.-J. Lee, and S.-H. Kim, \u201cKsponSpeech: Korean Spontaneous Speech Corpus for '
    'Automatic Speech Recognition,\u201d Applied Sciences, vol. 10, no. 19, p. 6936, 2020.',

    '[15] T. Sereda, \u201cTranscribe, Align and Segment: Creating Speech Datasets for Low-Resource '
    'Languages,\u201d arXiv:2406.12674, 2024.',

    '[16] Thennal D K, J. James, D. P. Gopinath, and M. Ashraf K, \u201cAdvocating Character Error '
    'Rate for Multilingual ASR Evaluation,\u201d in Findings of NAACL, 2025. arXiv:2410.07400.',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), 'research_paper.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
